#!/usr/bin/env python3
"""
Generate a simple ACM-like .docx from the proposal Markdown.
Requires: python-docx
Usage: python3 tools/generate_acm_docx.py qwen_proposal.md qwen_proposal.docx
"""
import sys
from docx import Document

MD_PATH = sys.argv[1] if len(sys.argv) > 1 else 'qwen_proposal.md'
OUT_PATH = sys.argv[2] if len(sys.argv) > 2 else 'qwen_proposal.docx'

def md_to_docx(md_path, out_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        s = line.rstrip('\n')
        if not s:
            doc.add_paragraph('')
            continue
        # Simple heuristics for headings and bullets
        if s.startswith('Title:'):
            title = s[len('Title:'):].strip()
            doc.add_heading(title, level=0)
            continue
        if s.startswith('Abstract'):
            doc.add_heading('Abstract', level=1)
            continue
        if s.startswith('Background') or s.startswith('Project Goals') or s.startswith('Planned Work') or s.startswith('Methodology') or s.startswith('References') or s.startswith('Repository Context') or s.startswith('Feasibility') or s.startswith('Next Steps'):
            # map broad headings
            doc.add_heading(s.strip(), level=2)
            continue
        if s.startswith('- '):
            doc.add_paragraph(s[2:].strip(), style='List Bullet')
            continue
        if s.startswith('+- '):
            doc.add_paragraph(s[3:].strip(), style='List Bullet')
            continue
        if s.startswith('1)') or s[0:2].isdigit() and s[2] in ').':
            doc.add_paragraph(s)
            continue
        # Normal paragraph
        doc.add_paragraph(s)

    doc.save(out_path)
    print(f'Wrote {out_path}')

if __name__ == '__main__':
    md_to_docx(MD_PATH, OUT_PATH)
